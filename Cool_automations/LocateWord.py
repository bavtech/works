from logging import raiseExceptions
import docx  
import os 
import pathlib 
import time 
workingDirectory = r'C:\Users\white\Documents\DAILY'
cv = r'C:\Users\white\Documents\DAILY\Others\CV_Temp'
keepON =  True 
currentYear  = time.localtime().tm_year 
aa =  pathlib.Path.cwd().rglob(f'*.docx') 
# TypesSupported = ['*.doc*', ]

def yearlySearch(year, searchPrase):

    # if year.isdigit():
    #     os.chdir(workingDirectory + f'\{year}')
    # elif year == '':
    
    if year.isdigit() and (int(year) >=2020  and int(year)<=currentYear) :
        
        os.chdir(workingDirectory + f'\{year}')

    elif year == '':

        os.chdir(workingDirectory ) 
        
    else:
        
        os.chdir(workingDirectory + f'\{currentYear}')
        
    

    all_files = pathlib.Path.cwd().rglob(f'*.docx')

    for File in all_files:
        if File.name.startswith('~$'):
                pass
        else:

            doc   = docx.Document(File)

            all_paragraphs =  doc.paragraphs

            for content in all_paragraphs:
                if searchPrase.lower() in  content.text.lower() :
                    print("found it")
                    print(File.name)
                    print(f"located at {File}")
                    print()
                    break


def cv_Search(searchPrase):

    os.chdir(cv)
    
    all_files = pathlib.Path.cwd().rglob(f'*.docx') 
    
    for File in all_files:
        
        try:
            if File.name.startswith('~$'):
                pass
            else:

                doc   = docx.Document(File)

                # print(f'now in {File}')
                # print()
                all_paragraphs =  doc.paragraphs
                
                for content in all_paragraphs:
                    if searchPrase.lower() in  content.text.lower().strip() :
                        print("found it")
                        print(f"File found\n file name: {File.name}")
                        input('Press enter to continue search after first spot')
        except ValueError as e:
            print(f"{File} is a .doc file and not a docx file and hence cannot be opened, moving to next file")

        


if __name__ == '__main__':

    
    while True:
     
        _ = int(input("To search for a word in CV press 1\n to Search for a word in a general file press 2"))

        if _ == 1:
            val = input("enter the word you wish to search for >> ")
            cv_Search(val)
        elif _ == 2:
            print("Please enter some details below ")
            print()
            val =  input("enter the word you wish to search for >> ")
            year = input("What year what year do you wish to search, press enter to pick current year as default  or enter a valid year ? >> ").strip()
            

           

            yearlySearch(year, val) 
        

